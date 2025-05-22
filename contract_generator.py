import os
import re
import fitz  # for PDF text extraction
from docx import Document
from docxtpl import DocxTemplate
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from utils.openai_client import generate_output  # GPT client for missing clauses

print("🚀 Script started!")

PLACEHOLDERS = [
    "Date", "Full_Name_Company_A", "Jurisdiction_A", "Address_A",
    "Full_Name_Company_B", "Jurisdiction_B", "Address_B",
    "Merger_Jurisdiction", "Share_Exchange_Ratio", "Cash_Per_Share",
    "Warranty_Survival_Years", "Closing_Date", "Governing_Law",
    "Arbitration_Organization", "Arbitration_Location",
    "Signatory_Name_A", "Signatory_Title_A", "Signatory_Name_B", "Signatory_Title_B", "Number_of_years"
]

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    return "\n".join(page.get_text() for page in doc)

def extract_text_from_docx(output_path):
    doc = Document(output_path)
    return "\n".join(para.text for para in doc.paragraphs)

def read_document_auto(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        print(f"📄 Detected PDF document: {path}")
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        print(f"📄 Detected DOCX document: {path}")
        return extract_text_from_docx(path)
    else:
        raise ValueError("Unsupported file format. Please provide .pdf or .docx")

def detect_contract_type(text):
    text_lower = text.lower()
    if "acquisition" in text_lower and "merger" not in text_lower:
        return "acquisition"
    elif "merger" in text_lower and "acquisition" not in text_lower:
        return "merger"
    elif "merger and acquisition" in text_lower or ("merger" in text_lower and "acquisition" in text_lower):
        return "merger"
    return "generic"
def extract_context_from_text(text):
    context = {key: "" for key in PLACEHOLDERS}
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    in_disclosing_party = False
    in_receiving_party = False

    for i, line in enumerate(lines):
        if "Disclosing Party" in line:
            in_disclosing_party = True
            in_receiving_party = False
            continue
        elif "Receiving Party" in line:
            in_receiving_party = True
            in_disclosing_party = False
            continue

        if in_disclosing_party:
            if line.lower().startswith("name:"):
                context["Full_Name_Company_A"] = line.split(":", 1)[1].strip()
            elif line.lower().startswith("jurisdiction:"):
                context["Jurisdiction_A"] = line.split(":", 1)[1].strip()
            elif line.lower().startswith("address:"):
                context["Address_A"] = line.split(":", 1)[1].strip()

        if in_receiving_party:
            if line.lower().startswith("name:"):
                context["Full_Name_Company_B"] = line.split(":", 1)[1].strip()
            elif line.lower().startswith("jurisdiction:"):
                context["Jurisdiction_B"] = line.split(":", 1)[1].strip()
            elif line.lower().startswith("address:"):
                context["Address_B"] = line.split(":", 1)[1].strip()

        if not context["Date"] and "effective date" in line.lower():
            context["Date"] = line.split("effective date")[-1].strip().strip(".")

        if not context["Closing_Date"] and "Closing shall take place on" in line:
            context["Closing_Date"] = line.split("on")[-1].strip().strip(".")

        if not context["Governing_Law"] and "laws of" in line:
            context["Governing_Law"] = line.split("laws of")[-1].strip().strip(".")

        if not context["Cash_Per_Share"] and "cash payment of" in line:
            context["Cash_Per_Share"] = line.split("cash payment of")[-1].strip().split()[0]
        
        if not context["Number_of_years"] and "Closing for a period of years" in line:
            context["Number_of_years"] = line.split("Closing for a period of years")[-1].strip().split()[0]

        if not context["Share_Exchange_Ratio"] and "receive" in line and "shares" in line:
            match = re.search(r"receive (\d+) shares", line)
            if match:
                context["Share_Exchange_Ratio"] = match.group(1)

        if not context["Warranty_Survival_Years"] and "survive the Closing" in line:
            match = re.search(r"period of (\d+)", line)
            if match:
                context["Warranty_Survival_Years"] = match.group(1)

        if "rules of" in line and not context["Arbitration_Organization"]:
            context["Arbitration_Organization"] = line.split("rules of")[-1].strip().strip(".")

        if "arbitration shall take place in" in line and not context["Arbitration_Location"]:
            context["Arbitration_Location"] = line.split("place in")[-1].strip().strip(".")

        if "Name:" in line:
            if not context["Signatory_Name_A"]:
                context["Signatory_Name_A"] = line.split("Name:")[-1].strip()
            elif not context["Signatory_Name_B"]:
                context["Signatory_Name_B"] = line.split("Name:")[-1].strip()

        if "Title:" in line:
            if not context["Signatory_Title_A"]:
                context["Signatory_Title_A"] = line.split("Title:")[-1].strip()
            elif not context["Signatory_Title_B"]:
                context["Signatory_Title_B"] = line.split("Title:")[-1].strip()

    # Address & Jurisdiction inference
    context["Merger_Jurisdiction"] = context["Governing_Law"]
    context["Address_A"] = "Address for Party A extracted manually if format varies"
    context["Address_B"] = "Address for Party B extracted manually if format varies"
    context["Jurisdiction_A"] = context["Governing_Law"]
    context["Jurisdiction_B"] = context["Governing_Law"]

    return context
# def render_contract_with_jinja2(data, templates_dir, output_html_path):
#     print("🌿 Rendering contract with Jinja2 template...")
#     env = Environment(loader=FileSystemLoader(templates_dir))
#     template = env.get_template("contract_template.jinja2")
#     rendered_content = template.render(data)

#     os.makedirs(os.path.dirname(output_html_path), exist_ok=True)
#     with open(output_html_path, "w", encoding="utf-8") as f:
#         f.write(rendered_content)
#     print(f"✅ Contract HTML generated at: {output_html_path}")
#     return rendered_content

# def convert_html_to_pdf(html_path, output_pdf_path):
#     print("📄 Converting HTML to PDF...")
#     os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
#     HTML(html_path).write_pdf(output_pdf_path)
#     print(f"✅ Contract PDF generated at: {output_pdf_path}")

def compare_and_add_missing_clauses(original_text, rendered_text, docx_path):
    print("🧠 Comparing with OpenAI for missing clauses...")

    prompt = f"""
You are a legal expert assistant.

Compare the following two contract texts:

1. Extracted from input document:
\"\"\"
{original_text}
\"\"\"

2. Rendered contract from template:
\"\"\"
{rendered_text}
\"\"\"

Identify clauses that exist in the input document but are missing in the template-generated version.
Return only the formal legal version of the missing clauses as a bulleted list of full clauses.
"""

    missing_clauses = generate_output(prompt)
    if missing_clauses:
        print("✍️ Appending missing clauses to DOCX...")
        doc = Document(docx_path)
        doc.add_paragraph("\n\nAdditional Clauses Identified from Source Document:\n")
        doc.add_paragraph(missing_clauses)
        doc.save(docx_path)
    else:
        print("✅ No missing clauses identified by OpenAI.")

def generate_contract_from_structured_doc(structured_doc_path, templates_dir, output_contract_path):
    print("📄 Reading structured document...")
    text = read_document_auto(structured_doc_path)

    print("🔍 Extracting structured context from document...")
    data = extract_context_from_text(text)
    print("📦 Extracted context data:", data)

    print("📌 Detecting contract type...")
    contract_type = detect_contract_type(text)
    print("🔎 Detected contract type:", contract_type)

    template_map = {
        "merger": os.path.join(templates_dir, "merger_template.docx"),
        "acquisition": os.path.join(templates_dir, "acquisition_template.docx"),
        "generic": os.path.join(templates_dir, "generic_template.docx")
    }

    template_path = template_map.get(contract_type)
    print("📂 Using template:", template_path)

    if not os.path.exists(template_path):
        print(f"❌ Template file not found: {template_path}")
        return

    print("🧩 Rendering contract (DOCX)...")
    doc = DocxTemplate(template_path)
    doc.render(data)
    os.makedirs(os.path.dirname(output_contract_path), exist_ok=True)
    doc.save(output_contract_path)

    # Extract rendered text for GPT comparison
    rendered_text = extract_text_from_docx(output_contract_path)

    # Use OpenAI to find & add missing clauses
    compare_and_add_missing_clauses(original_text=text, rendered_text=rendered_text, docx_path=output_contract_path)
    print(f"✅ Final contract (with added clauses) saved at: {output_contract_path}")

    # # Generate HTML + PDF
    # output_html_path = output_contract_path.replace(".docx", ".html")
    # render_contract_with_jinja2(data, templates_dir, output_html_path)

    # output_pdf_path = output_contract_path.replace(".docx", ".pdf")
    # convert_html_to_pdf(output_html_path, output_pdf_path)

if __name__ == "__main__":
    generate_contract_from_structured_doc(
        structured_doc_path="input_docs/sample_input.docx",  # or .pdf
        templates_dir="templates",
        output_contract_path="output_contracts/generated_contract.docx"
    )
