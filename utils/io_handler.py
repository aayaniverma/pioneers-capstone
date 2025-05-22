from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def write_docx(content, output_path):
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)