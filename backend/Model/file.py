# import pandas as pd
# import json
# import os
# from docx import Document
# from tqdm import tqdm

# DOC_FOLDER = "data/docs"
# CSV_PATH = "data/label.csv"
# OUTPUT_PATH = "qa_dataset.json"

# fields = [
#     "Date", "Full_Name_Company_A", "Jurisdiction_A", "Address_A",
#     "Full_Name_Company_B", "Jurisdiction_B", "Address_B",
#     "Merger_Jurisdiction", "Share_Exchange_Ratio", "Cash_Per_Share",
#     "Warranty_Survival_Years", "Closing_Date", "Governing_Law",
#     "Arbitration_Organization", "Arbitration_Location",
#     "Signatory_Name_A", "Signatory_Title_A", "Signatory_Name_B",
#     "Signatory_Title_B", "Number_of_years"
# ]

# df = pd.read_csv(CSV_PATH)
# qa_data = {"data": []}

# for _, row in tqdm(df.iterrows(), total=len(df)):
#     filename = row['filename']
#     if not filename.endswith('.docx'):
#         filename += '.docx'

#     full_path = os.path.join(DOC_FOLDER, filename)
#     if not os.path.exists(full_path):
#         print(f"⚠️ File not found: {full_path}")
#         continue

#     doc = Document(full_path)
#     context = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

#     for field in fields:
#         answer = row.get(field)
#         if pd.isna(answer) or not answer.strip():
#             continue  # skip missing labels

#         qa_data["data"].append({
#             "context": context,
#             "question": f"What is the {field}?",
#             "answers": [{"text": answer.strip()}]
#         })

# with open(OUTPUT_PATH, "w") as f:
#     json.dump(qa_data, f, indent=2)

# print(f"✅ QA dataset saved to {OUTPUT_PATH}")


import pandas as pd
import json
import os
from docx import Document
from tqdm import tqdm

DOC_FOLDER = "data/docs"
CSV_PATH = "data/label.csv"
OUTPUT_PATH = "qa_dataset.json"

fields = [
    "Date", "Full_Name_Company_A", "Jurisdiction_A", "Address_A",
    "Full_Name_Company_B", "Jurisdiction_B", "Address_B",
    "Merger_Jurisdiction", "Share_Exchange_Ratio", "Cash_Per_Share",
    "Warranty_Survival_Years", "Closing_Date", "Governing_Law",
    "Arbitration_Organization", "Arbitration_Location",
    "Signatory_Name_A", "Signatory_Title_A", "Signatory_Name_B",
    "Signatory_Title_B", "Number_of_years"
]

df = pd.read_csv(CSV_PATH)
qa_data = {"data": []}

for _, row in tqdm(df.iterrows(), total=len(df)):
    filename = row['filename']
    if not filename.endswith('.docx'):
        filename += '.docx'

    full_path = os.path.join(DOC_FOLDER, filename)
    if not os.path.exists(full_path):
        print(f"⚠️ File not found: {full_path}")
        continue

    doc = Document(full_path)
    context = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    for field in fields:
        answer = row.get(field)
        if pd.isna(answer) or not str(answer).strip():
            continue

        answer = str(answer).strip()
        answer_start = context.find(answer)

        if answer_start == -1:
            print(f"⚠️ '{answer}' not found in document: {filename}")
            continue

        qa_data["data"].append({
            "context": context,
            "question": f"What is the {field.replace('_', ' ')}?",
            "answers": {
                "text": [answer],
                "answer_start": [answer_start]
            }
        })

with open(OUTPUT_PATH, "w") as f:
    json.dump(qa_data, f, indent=2)

print(f"✅ QA dataset saved to {OUTPUT_PATH}")