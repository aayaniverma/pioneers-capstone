# import os
# import json
# import torch
# from docx import Document
# from pathlib import Path
# from transformers import AutoTokenizer, AutoModelForQuestionAnswering
# import re
# from tqdm import tqdm

# # Load model and tokenizer
# model_path = "contract_qa_model"
# tokenizer = AutoTokenizer.from_pretrained(model_path)
# model = AutoModelForQuestionAnswering.from_pretrained(model_path)
# model.eval()

# # Questions
# questions = [
#     "What is the Date?",
#     "What is the Full Name Company A?",
#     "What is the Jurisdiction A?",
#     "What is the Address A?",
#     "What is the Full Name Company B?",
#     "What is the Jurisdiction B?",
#     "What is the Address B?",
#     "What is the Merger Jurisdiction?",
#     "What is the Share Exchange Ratio?",
#     "What is the Cash Per Share?",
#     "What is the Warranty Survival Years?",
#     "What is the Closing Date?",
#     "What is the Governing Law?",
#     "What is the Arbitration Organization?",
#     "What is the Arbitration Location?",
#     "What is the Signatory Name A?",
#     "What is the Signatory Title A?",
#     "What is the Signatory Name B?",
#     "What is the Signatory Title B?",
#     "What is the Number of years?"
# ]

# # Constants
# MAX_ANSWER_LEN = 30
# CONFIDENCE_THRESHOLD = 0.5
# test_docs_path = Path("data/test_docs")

# # Extract text from .docx
# def extract_text(doc_path):
#     doc = Document(doc_path)
#     return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

# # Extract signatory data with regex fallback
# def extract_signatories(text):
#     sigs = {
#         "What is the Signatory Name A?": "",
#         "What is the Signatory Title A?": "",
#         "What is the Signatory Name B?": "",
#         "What is the Signatory Title B?": ""
#     }
#     try:
#         pattern = r"Signed by authorized representatives:\s*- (.*?), (.*?), for .*?\n- (.*?), (.*?), for"
#         match = re.search(pattern, text, re.DOTALL)
#         if match:
#             sigs["What is the Signatory Name A?"] = match.group(1).strip()
#             sigs["What is the Signatory Title A?"] = match.group(2).strip()
#             sigs["What is the Signatory Name B?"] = match.group(3).strip()
#             sigs["What is the Signatory Title B?"] = match.group(4).strip()
#     except Exception as e:
#         pass
#     return sigs

# # Prediction
# def predict_answer(question, context):
#     inputs = tokenizer(question, context, return_tensors="pt", truncation=True, max_length=512)
#     with torch.no_grad():
#         outputs = model(**inputs)

#     start_logits = outputs.start_logits
#     end_logits = outputs.end_logits
#     start_probs = torch.softmax(start_logits, dim=1)
#     end_probs = torch.softmax(end_logits, dim=1)

#     max_start = torch.max(start_probs).item()
#     max_end = torch.max(end_probs).item()

#     if max_start < CONFIDENCE_THRESHOLD or max_end < CONFIDENCE_THRESHOLD:
#         return ""

#     start_idx = torch.argmax(start_logits).item()
#     end_idx = torch.argmax(end_logits).item() + 1

#     if end_idx - start_idx > MAX_ANSWER_LEN:
#         end_idx = start_idx + MAX_ANSWER_LEN

#     answer = tokenizer.decode(inputs["input_ids"][0][start_idx:end_idx], skip_special_tokens=True)
#     return answer.strip()

# # Run for all test documents
# results = {}
# for file_path in tqdm(test_docs_path.glob("*.docx")):
#     context = extract_text(file_path)
#     sig_fallbacks = extract_signatories(context)
#     qa_output = {}

#     for question in questions:
#         if question in sig_fallbacks and sig_fallbacks[question]:
#             qa_output[question] = sig_fallbacks[question]
#         else:
#             qa_output[question] = predict_answer(question, context)

#     results[file_path.name] = qa_output

# # Save results
# output_file = "predictions.json"
# with open(output_file, "w") as f:
#     json.dump(results, f, indent=2)

# print(f"✅ Predictions saved to {output_file}")

import os
import json
import torch
from docx import Document
from pathlib import Path
from transformers import AutoTokenizer, AutoModelForQuestionAnswering
import re
from tqdm import tqdm
from Model.contract_generator import generate_contract

# Load model and tokenizer
model_path = os.path.join(os.path.dirname(__file__), "contract_qa_model")

tokenizer = AutoTokenizer.from_pretrained(model_path)
model = AutoModelForQuestionAnswering.from_pretrained(model_path)
model.eval()

# Questions
questions = [
    "What is the Date?",
    "What is the Full Name Company A?",
    "What is the Jurisdiction A?",
    "What is the Address A?",
    "What is the Full Name Company B?",
    "What is the Jurisdiction B?",
    "What is the Address B?",
    "What is the Merger Jurisdiction?",
    "What is the Share Exchange Ratio?",
    "What is the Cash Per Share?",
    "What is the Warranty Survival Years?",
    "What is the Closing Date?",
    "What is the Governing Law?",
    "What is the Arbitration Organization?",
    "What is the Arbitration Location?",
    "What is the Signatory Name A?",
    "What is the Signatory Title A?",
    "What is the Signatory Name B?",
    "What is the Signatory Title B?",
    "What is the Number of years?"
]

# Constants
MAX_ANSWER_LEN = 30
CONFIDENCE_THRESHOLD = 0.5
test_docs_path = Path("../temp/tempdoc")

# Extract text from .docx
def extract_text(doc_path):
    doc = Document(doc_path)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

# Extract signatory data with regex fallback
def extract_signatories(text):
    sigs = {
        "What is the Signatory Name A?": "",
        "What is the Signatory Title A?": "",
        "What is the Signatory Name B?": "",
        "What is the Signatory Title B?": ""
    }
    try:
        pattern = r"Signed by authorized representatives:\s*- (.*?), (.*?), for .*?\n- (.*?), (.*?), for"
        match = re.search(pattern, text, re.DOTALL)
        if match:
            sigs["What is the Signatory Name A?"] = match.group(1).strip()
            sigs["What is the Signatory Title A?"] = match.group(2).strip()
            sigs["What is the Signatory Name B?"] = match.group(3).strip()
            sigs["What is the Signatory Title B?"] = match.group(4).strip()
    except Exception:
        pass
    return sigs

# Prediction
def predict_answer(question, context):
    inputs = tokenizer(question, context, return_tensors="pt", truncation=True, max_length=512)
    with torch.no_grad():
        outputs = model(**inputs)

    start_logits = outputs.start_logits
    end_logits = outputs.end_logits
    start_probs = torch.softmax(start_logits, dim=1)
    end_probs = torch.softmax(end_logits, dim=1)

    max_start = torch.max(start_probs).item()
    max_end = torch.max(end_probs).item()

    if max_start < CONFIDENCE_THRESHOLD or max_end < CONFIDENCE_THRESHOLD:
        return ""

    start_idx = torch.argmax(start_logits).item()
    end_idx = torch.argmax(end_logits).item() + 1

    if end_idx - start_idx > MAX_ANSWER_LEN:
        end_idx = start_idx + MAX_ANSWER_LEN

    answer = tokenizer.decode(inputs["input_ids"][0][start_idx:end_idx], skip_special_tokens=True)
    answer = answer.strip()

    # Heuristic: trim at newline or long trailing artifacts
    if any(sep in answer for sep in [". ", "\n", "\r"]):
        answer = re.split(r"[\n\r\.]", answer)[0].strip()

    return answer

# Run for all test documents
def run_model_pipeline(doc_path):
    context = extract_text(doc_path)
    sig_fallbacks = extract_signatories(context)
    qa_output = {}

    for question in questions:
        if question in sig_fallbacks and sig_fallbacks[question]:
            qa_output[question] = sig_fallbacks[question]
        else:
            qa_output[question] = predict_answer(question, context)

    predictions = {os.path.basename(doc_path): qa_output}
    output_path = generate_contract(predictions)
    print(f"Returning generated path: {output_path}")

    return output_path



