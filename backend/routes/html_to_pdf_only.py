from fastapi import APIRouter, Form
from fastapi.responses import JSONResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
from docx2pdf import convert
import os

router = APIRouter()
TEMP_PRE_DIR = "temp/temppre"
os.makedirs(TEMP_PRE_DIR, exist_ok=True)

@router.post("/html_to_pdf_only/")
async def html_to_pdf_only(
    user_text: str = Form(...),
    filename: str = Form("output.docx")
):
    soup = BeautifulSoup(user_text, "html.parser")
    doc = Document()

    # Set base font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Handle common tags
    for elem in soup.find_all(['h1', 'h2', 'p', 'b', 'i', 'ul', 'ol', 'li']):
        if elem.name == 'h1':
            p = doc.add_paragraph(elem.get_text(strip=True), style='Heading 1')
        elif elem.name == 'h2':
            p = doc.add_paragraph(elem.get_text(strip=True), style='Heading 2')
        elif elem.name == 'p':
            p = doc.add_paragraph(elem.get_text(strip=True), style='Normal')
        elif elem.name == 'ul':
            for li in elem.find_all('li'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(li.get_text(strip=True))
        elif elem.name == 'ol':
            for li in elem.find_all('li'):
                p = doc.add_paragraph(style='List Number')
                p.add_run(li.get_text(strip=True))
        elif elem.name == 'b':
            p = doc.add_paragraph()
            run = p.add_run(elem.get_text(strip=True))
            run.bold = True
        elif elem.name == 'i':
            p = doc.add_paragraph()
            run = p.add_run(elem.get_text(strip=True))
            run.italic = True

    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    docx_path = os.path.join(TEMP_PRE_DIR, filename)
    pdf_path = docx_path.replace(".docx", ".pdf")

    doc.save(docx_path)

    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        return JSONResponse(content={"error": f"PDF conversion failed: {str(e)}"}, status_code=500)

    os.remove(docx_path)

    return JSONResponse(
        content={
            "message": f"PDF saved as {os.path.basename(pdf_path)}.",
            "file_path": pdf_path
        },
        status_code=200
    )
