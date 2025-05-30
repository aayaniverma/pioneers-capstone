from fastapi import APIRouter, UploadFile, File, HTTPException
import shutil
import os
import re
#import fitz  # for PDF text extraction
from docx import Document
from docxtpl import DocxTemplate
from jinja2 import Environment, FileSystemLoader
#from weasyprint import HTML
from utils.openai_client import generate_output  # GPT client for missing clauses

router = APIRouter()
@router.post("/generate-contract/")
async def generate_contract(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are accepted")

    # Save uploaded file temporarily to output_docs instead of temp_uploads
    os.makedirs("output_docs", exist_ok=True)
    temp_input_path = f"output_docs/{file.filename}"
    with open(temp_input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    output_path = "output_contracts/generated_contract.docx"
    templates_dir = "templates"

    try:
        generate_contract_from_structured_doc(
            structured_doc_path=temp_input_path,
            templates_dir=templates_dir,
            output_contract_path=output_path,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Contract generation failed: {e}")

    return {"message": "Contract generated successfully", "output_path": output_path}


print("🚀 Script started!")

PLACEHOLDERS = [
    "Date", "Full_Name_Company_A", "Jurisdiction_A", "Address_A",
    "Full_Name_Company_B", "Jurisdiction_B", "Address_B",
    "Merger_Jurisdiction", "Share_Exchange_Ratio", "Cash_Per_Share",
    "Warranty_Survival_Years", "Closing_Date", "Governing_Law",
    "Arbitration_Organization", "Arbitration_Location",
    "Signatory_Name_A", "Signatory_Title_A", "Signatory_Name_B", "Signatory_Title_B", "Number_of_years"
]
COMPANY_SUFFIXES = ['Inc.', 'Ltd.', 'LLC', 'Pvt. Ltd.', 'Corporation', 'Corp.', 'Limited', 'Company']
ADDRESS_PATTERNS = ['street', 'st.', 'road', 'rd.', 'avenue', 'ave', 'lane', 'ln.', 'drive', 'dr.', 'boulevard', 'blvd', 'floor', 'fl.', 'suite', 'building', 'bldg']

def looks_like_company_name(line):
    return any(suffix.lower() in line.lower() for suffix in COMPANY_SUFFIXES)

def looks_like_address(line):
    return any(p in line.lower() for p in ADDRESS_PATTERNS) or bool(re.search(r'\d{5}|\b[A-Z]{2,3}\b', line))

def clean_extracted_value(val):
    # Remove common prefixes, trailing commas and spaces
    val = val.strip()
    for prefix in ["Name:", "name:", "Address:", "address:"]:
        if val.lower().startswith(prefix.lower()):
            val = val[len(prefix):].strip()
    val = val.strip(", ").strip()
    return val

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    return "\n".join(page.get_text() for page in doc)

def extract_text_from_docx(output_path):
    doc = Document(output_path)
    return "\n".join(para.text for para in doc.paragraphs)

def read_document_auto(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        print(f"📄 Detected PDF document: {path}")
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        print(f"📄 Detected DOCX document: {path}")
        return extract_text_from_docx(path)
    else:
        raise ValueError("Unsupported file format. Please provide .pdf or .docx")

def detect_contract_type(text):
    text_lower = text.lower()
    if "acquisition" in text_lower and "merger" not in text_lower:
        return "acquisition"
    elif "merger" in text_lower and "acquisition" not in text_lower:
        return "merger"
    elif "merger and acquisition" in text_lower or ("merger" in text_lower and "acquisition" in text_lower):
        return "merger"
    return "generic"
def extract_context_from_text(text):
    context = {key: "" for key in PLACEHOLDERS}
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    in_disclosing_party = False
    in_receiving_party = False

    for line in lines:
        line_lower = line.lower()

        if "disclosing party" in line_lower:
            in_disclosing_party = True
            in_receiving_party = False
            continue
        elif "receiving party" in line_lower:
            in_receiving_party = True
            in_disclosing_party = False
            continue

        # Company name (label-based or pattern-based)
        if ("name:" in line_lower or looks_like_company_name(line)) and not ("signatory" in line_lower):
            value = re.sub(r'(?i)^name\s*:\s*', '', line).strip()
            value = clean_extracted_value(value)
            if in_disclosing_party and not context["Full_Name_Company_A"]:
                context["Full_Name_Company_A"] = value
            elif in_receiving_party and not context["Full_Name_Company_B"]:
                context["Full_Name_Company_B"] = value

        # Jurisdiction
        if "jurisdiction:" in line_lower:
            value = line.split(":", 1)[1].strip()
            if in_disclosing_party and not context["Jurisdiction_A"]:
                context["Jurisdiction_A"] = value
            elif in_receiving_party and not context["Jurisdiction_B"]:
                context["Jurisdiction_B"] = value

        # Address (label-based or pattern-based)
        if "address:" in line_lower or looks_like_address(line):
            value = re.sub(r'(?i)^address\s*:\s*', '', line).strip()
            value = clean_extracted_value(value)
            if in_disclosing_party and not context["Address_A"]:
                context["Address_A"] = value
            elif in_receiving_party and not context["Address_B"]:
                context["Address_B"] = value
        # Other fields
        if not context["Date"] and "effective as of" in line_lower:
            match = re.search(r"effective as of (.+)", line_lower)
            if match:
                context["Date"] = match.group(1).strip().rstrip(".")

        if not context["Governing_Law"] and "laws of" in line_lower:
            context["Governing_Law"] = line.split("laws of")[-1].strip().rstrip(".")

        if not context["Cash_Per_Share"] and "cash payment of" in line_lower:
            context["Cash_Per_Share"] = line.split("cash payment of")[-1].strip().split()[0]

        if not context["Closing_Date"] and "closing shall take place on" in line_lower:
            context["Closing_Date"] = line.split("on")[-1].strip().strip(".")

        if not context["Share_Exchange_Ratio"] and "receive" in line_lower and "shares" in line_lower:
            match = re.search(r"receive (\d+) shares", line_lower)
            if match:
                context["Share_Exchange_Ratio"] = match.group(1)

        if not context["Warranty_Survival_Years"] and "survive the closing" in line_lower:
            match = re.search(r"period of (\d+)", line_lower)
            if match:
                context["Warranty_Survival_Years"] = match.group(1)

        if not context["Number_of_years"] and "closing for a period of years" in line_lower:
            match = re.search(r"closing for a period of years (\d+)", line_lower)
            if match:
                context["Number_of_years"] = match.group(1)

        if "rules of" in line_lower and not context["Arbitration_Organization"]:
            context["Arbitration_Organization"] = line.split("rules of")[-1].strip().strip(".")

        if "arbitration shall take place in" in line_lower and not context["Arbitration_Location"]:
            context["Arbitration_Location"] = line.split("place in")[-1].strip().strip(".")

        if "name:" in line_lower:
            if not context["Signatory_Name_A"]:
                context["Signatory_Name_A"] = line.split(":", 1)[1].strip()
            elif not context["Signatory_Name_B"]:
                context["Signatory_Name_B"] = line.split(":", 1)[1].strip()

        if "title:" in line_lower:
            if not context["Signatory_Title_A"]:
                context["Signatory_Title_A"] = line.split(":", 1)[1].strip()
            elif not context["Signatory_Title_B"]:
                context["Signatory_Title_B"] = line.split(":", 1)[1].strip()

    # Set merger jurisdiction same as governing law if not otherwise extracted
    context["Merger_Jurisdiction"] = context["Governing_Law"]

    return context
# def render_contract_with_jinja2(data, templates_dir, output_html_path):
#     print("🌿 Rendering contract with Jinja2 template...")
#     env = Environment(loader=FileSystemLoader(templates_dir))
#     template = env.get_template("contract_template.jinja2")
#     rendered_content = template.render(data)

#     os.makedirs(os.path.dirname(output_html_path), exist_ok=True)
#     with open(output_html_path, "w", encoding="utf-8") as f:
#         f.write(rendered_content)
#     print(f"✅ Contract HTML generated at: {output_html_path}")
#     return rendered_content

# def convert_html_to_pdf(html_path, output_pdf_path):
#     print("📄 Converting HTML to PDF...")
#     os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
#     HTML(html_path).write_pdf(output_pdf_path)
#     print(f"✅ Contract PDF generated at: {output_pdf_path}")

def compare_and_add_missing_clauses(original_text, rendered_text, docx_path):
    print("🧠 Comparing with OpenAI for missing clauses...")

    prompt = f"""
You are a legal expert assistant.

Compare the following two contract texts:

1. Extracted from input document:
\"\"\"
{original_text}
\"\"\"

2. Rendered contract from template:
\"\"\"
{rendered_text}
\"\"\"

Identify clauses that exist in the input document but are missing in the template-generated version.
Return only the formal legal version of the missing clauses as a bulleted list of full clauses.
"""

    missing_clauses = generate_output(prompt)
    if missing_clauses:
        print("✍️ Appending missing clauses to DOCX...")
        doc = Document(docx_path)
        doc.add_paragraph("\n\nAdditional Clauses Identified from Source Document:\n")
        doc.add_paragraph(missing_clauses)
        doc.save(docx_path)
    else:
        print("✅ No missing clauses identified by OpenAI.")

def generate_contract_from_structured_doc(structured_doc_path, templates_dir, output_contract_path):
    print("📄 Reading structured document...")
    text = read_document_auto(structured_doc_path)

    print("🔍 Extracting structured context from document...")
    data = extract_context_from_text(text)
    print("📦 Extracted context data:", data)

    print("📌 Detecting contract type...")
    contract_type = detect_contract_type(text)
    print("🔎 Detected contract type:", contract_type)

    template_map = {
        "merger": os.path.join(templates_dir, "merger_template.docx"),
        "acquisition": os.path.join(templates_dir, "acquisition_template.docx"),
        "generic": os.path.join(templates_dir, "generic_template.docx")
    }

    template_path = template_map.get(contract_type)
    print("📂 Using template:", template_path)

    if not os.path.exists(template_path):
        print(f"❌ Template file not found: {template_path}")
        return

    print("🧩 Rendering contract (DOCX)...")
    doc = DocxTemplate(template_path)
    def sanitize_value(val):
        if not val or not val.strip():
            return "N/A"
        return val.replace("\n", " ").strip()

    safe_data = {k: sanitize_value(v) for k, v in data.items()}
    empty_fields = [k for k, v in data.items() if not v.strip()]
    if empty_fields:
        print(f"⚠️ Warning: The following fields were empty and replaced with 'N/A': {empty_fields}")
    safe_data = {k: (v.strip() if v.strip() else "N/A") for k, v in data.items()}

    try:
        doc.render(safe_data)
    except Exception as e:
        print(f"❌ Rendering failed: {e}")
        raise e
    os.makedirs(os.path.dirname(output_contract_path), exist_ok=True)
    doc.save(output_contract_path)

    # Extract rendered text for GPT comparison
    rendered_text = extract_text_from_docx(output_contract_path)

    # Use OpenAI to find & add missing clauses
    compare_and_add_missing_clauses(original_text=text, rendered_text=rendered_text, docx_path=output_contract_path)
    print(f"✅ Final contract (with added clauses) saved at: {output_contract_path}")

    # # Generate HTML + PDF
    # output_html_path = output_contract_path.replace(".docx", ".html")
    # render_contract_with_jinja2(data, templates_dir, output_html_path)

    # output_pdf_path = output_contract_path.replace(".docx", ".pdf")
    # convert_html_to_pdf(output_html_path, output_pdf_path)

#if __name__ == "__main__":
 #   generate_contract_from_structured_doc(
  #      structured_doc_path="output_docs/structured_output.docx",  # or .pdf
   #     templates_dir="templates",
    #    output_contract_path="output_contracts/generated_contract.docx"
    #)

    

 